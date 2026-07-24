#!/usr/bin/env python3
"""
将教学设计 Markdown 文件转换为格式化的 Word 文档。

用法：
    python3 to_docx.py input.md [output.docx]

如果不指定输出文件名，会自动使用 {课题}_教学设计.docx。

依赖：python-docx（已预装）
"""

import sys
import re
import subprocess
import tempfile
from pathlib import Path

try:
    from docx import Document
except ImportError:
    import subprocess
    subprocess.check_call(
        [sys.executable, '-m', 'pip', 'install', '-q', 'python-docx'],
        stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
    )
    from docx import Document

from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn


IMAGE_RE = re.compile(r'^!\[([^\]]*)\]\(([^)]+)\)$')


def set_chinese_font(run, font_name="宋体", size=Pt(12)):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = size
    # 设置中文字体
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def write_rich_text(paragraph, text, font_name="宋体", size=Pt(12)):
    """将含 **加粗** 标记的文本写入段落，正确渲染为 Word 加粗"""
    paragraph.clear()
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            set_chinese_font(run, font_name, size)
        else:
            run = paragraph.add_run(part)
            set_chinese_font(run, font_name, size)


# 表格单元格内需要分段的标签
_CELL_SECTION_MARKERS = re.compile(
    r'(?=\*\*(?:教师活动|学生活动|【设计意图】|活动[一二三四五六七八九十])[：:]?\*\*)'
)


def _clean_html_in_text(text):
    """清除 Markdown 表格中常见的 HTML 标签和实体。
    AI 在单行 Markdown 表格里经常用 <br> 换行、&nbsp; 做缩进。"""
    # <br> / <br/> / <br /> → 换行符
    text = re.sub(r'<br\s*/?>', '\n', text)
    # &nbsp; → 普通空格
    text = text.replace('&nbsp;', ' ')
    # 其他常见 HTML 实体
    text = text.replace('&lt;', '<')
    text = text.replace('&gt;', '>')
    text = text.replace('&amp;', '&')
    text = text.replace('&quot;', '"')
    return text


def write_cell_rich_text(cell, text, font_name="宋体", size=Pt(11)):
    """将表格单元格内容按 **教师活动：** / **学生活动：** / **【设计意图】** 等标签分段。
    先清理 HTML 标签，再按标签分段，每段在单元格里是独立段落。"""
    # 第一步：清理 HTML
    text = _clean_html_in_text(text)

    # 第二步：按标签位置分段
    sections = _CELL_SECTION_MARKERS.split(text)
    # 如果没有标签，就按换行符分段
    if len(sections) <= 1:
        sections = [s.strip() for s in text.split('\n') if s.strip()]
    else:
        # 标签分段后，每段内部的换行变成段内换行
        cleaned = []
        for s in sections:
            s = s.strip()
            if s:
                # 段内多个连续换行合并为一个空格
                s = re.sub(r'\n{2,}', '\n', s)
                cleaned.append(s)
        sections = cleaned

    if not sections:
        return

    # 第一段写入 cell 自带的段落
    first_para = cell.paragraphs[0]
    write_rich_text(first_para, sections[0].replace('\n', ' '), font_name, size)
    first_para.paragraph_format.space_after = Pt(2)
    first_para.paragraph_format.space_before = Pt(0)

    # 后续段落用 add_paragraph
    for section in sections[1:]:
        p = cell.add_paragraph()
        write_rich_text(p, section.replace('\n', ' '), font_name, size)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(4)


def parse_markdown(md_text):
    """解析 Markdown 文本为结构化数据"""
    lines = md_text.split('\n')
    elements = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # 代码块（竖式计算、板书示意等）
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                if lines[i].strip():
                    code_lines.append(lines[i])
                i += 1
            if code_lines:
                elements.append(('code', code_lines))
        # 标题
        elif line.startswith('# '):
            elements.append(('h1', line[2:].strip()))
        elif line.startswith('## '):
            elements.append(('h2', line[3:].strip()))
        elif line.startswith('### '):
            elements.append(('h3', line[4:].strip()))
        # 图片
        elif IMAGE_RE.match(line.strip()):
            match = IMAGE_RE.match(line.strip())
            elements.append(('image', (match.group(1).strip(), match.group(2).strip())))
        # 表格
        elif '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                if '---' not in lines[i]:
                    cells = [c.strip() for c in lines[i].split('|')[1:-1]]
                    table_lines.append(cells)
                i += 1
            elements.append(('table', table_lines))
            continue
        # 列表项
        elif line.startswith('- '):
            elements.append(('bullet', line[2:].strip()))
        # 加粗段落
        elif line.startswith('**') and line.endswith('**'):
            elements.append(('bold', line.strip('*').strip()))
        elif line.startswith('**') and '：**' in line:
            elements.append(('bold_label', line.strip()))
        # 分隔线
        elif line.strip() == '---':
            elements.append(('hr', ''))
        # 普通段落
        elif line.strip():
            elements.append(('para', line.strip()))

        i += 1

    return elements


def resolve_image_path(image_ref, base_dir):
    """解析 Markdown 图片路径，支持相对当前 Markdown 文件的路径。"""
    image_ref = image_ref.strip()
    if image_ref.startswith('<') and image_ref.endswith('>'):
        image_ref = image_ref[1:-1]
    path = Path(image_ref)
    if not path.is_absolute():
        path = (base_dir / path).resolve()
    return path


def convert_svg_to_png(svg_path):
    """将 SVG 转为临时 PNG，便于 python-docx 稳定嵌入 Word。
    优先使用 macOS sips，失败时 fallback 到 cairosvg（跨平台）。"""
    tmpdir = Path(tempfile.mkdtemp(prefix='svg-docx-'))
    output_png = tmpdir / f'{svg_path.stem}.png'

    # 方案一：macOS sips
    try:
        subprocess.run(
            ['sips', '-s', 'format', 'png', str(svg_path), '--out', str(output_png)],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        if output_png.exists():
            return output_png
    except (FileNotFoundError, subprocess.CalledProcessError):
        pass

    # 方案二：cairosvg（跨平台 fallback）
    try:
        import cairosvg
    except ImportError:
        subprocess.check_call(
            [sys.executable, '-m', 'pip', 'install', '-q', 'cairosvg'],
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
        )
        import cairosvg
    cairosvg.svg2png(url=str(svg_path), write_to=str(output_png))

    if not output_png.exists():
        raise RuntimeError(f"SVG 转 PNG 失败：{svg_path}")
    return output_png


def create_docx(elements, output_path, base_dir):
    """根据解析后的元素生成 Word 文档"""
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    BLACK = RGBColor(0x00, 0x00, 0x00)

    for elem_type, content in elements:
        if elem_type == 'h1':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(content)
            run.bold = True
            run.font.color.rgb = BLACK
            set_chinese_font(run, "黑体", Pt(22))

        elif elem_type == 'h2':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(content)
            run.bold = True
            run.font.color.rgb = BLACK
            set_chinese_font(run, "黑体", Pt(16))

        elif elem_type == 'h3':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(content)
            run.bold = True
            run.font.color.rgb = BLACK
            set_chinese_font(run, "黑体", Pt(14))

        elif elem_type == 'table':
            if len(content) >= 1:
                rows = len(content)
                cols = len(content[0]) if content else 0
                if cols > 0:
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    if cols == 4:
                        widths = [Cm(2.5), Cm(3.5), Cm(8.5), Cm(3.5)]
                        for i_row, row in enumerate(table.rows):
                            for j_col, cell in enumerate(row.cells):
                                cell.width = widths[j_col]
                    for i, row_data in enumerate(content):
                        for j, cell_text in enumerate(row_data):
                            if j < cols:
                                cell = table.rows[i].cells[j]
                                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                                # 清理 HTML 标签
                                clean_text = _clean_html_in_text(cell_text)
                                # 师生活动列（第3列）用分段渲染，其他列正常渲染
                                if i > 0 and cols == 4 and j == 2:
                                    write_cell_rich_text(cell, cell_text, "宋体", Pt(11))
                                else:
                                    write_rich_text(cell.paragraphs[0], clean_text.replace('\n', ' '), "宋体", Pt(11))
                                for paragraph in cell.paragraphs:
                                    paragraph.paragraph_format.space_after = Pt(2)
                    for j in range(cols):
                        for paragraph in table.rows[0].cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                set_chinese_font(run, "黑体", Pt(11))
                    doc.add_paragraph()  # 表格后空行

        elif elem_type == 'image':
            alt_text, image_ref = content
            image_path = resolve_image_path(image_ref, base_dir)
            if not image_path.exists():
                p = doc.add_paragraph()
                run = p.add_run(f"[图片缺失：{image_ref}]")
                run.bold = True
                set_chinese_font(run, "宋体", Pt(12))
                continue

            embed_path = image_path
            if image_path.suffix.lower() == '.svg':
                embed_path = convert_svg_to_png(image_path)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(embed_path), width=Cm(12.5))
            if alt_text:
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption.add_run(alt_text)
                caption_run.italic = True
                set_chinese_font(caption_run, "宋体", Pt(10))

        elif elem_type == 'bullet':
            # 处理加粗标记
            p = doc.add_paragraph(style='List Bullet')
            # 解析内联加粗
            parts = re.split(r'(\*\*.*?\*\*)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    set_chinese_font(run, "宋体", Pt(12))
                else:
                    run = p.add_run(part)
                    set_chinese_font(run, "宋体", Pt(12))

        elif elem_type == 'bold':
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.bold = True
            set_chinese_font(run, "黑体", Pt(12))

        elif elem_type == 'bold_label':
            p = doc.add_paragraph()
            # 解析 **标签：** 内容 格式
            parts = re.split(r'(\*\*.*?\*\*)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    set_chinese_font(run, "黑体", Pt(12))
                else:
                    run = p.add_run(part)
                    set_chinese_font(run, "宋体", Pt(12))

        elif elem_type == 'code':
            for code_line in content:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.left_indent = Cm(1.0)
                run = p.add_run(code_line)
                set_chinese_font(run, "Consolas", Pt(11))

        elif elem_type == 'para':
            p = doc.add_paragraph()
            # 解析内联加粗
            parts = re.split(r'(\*\*.*?\*\*)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    set_chinese_font(run, "宋体", Pt(12))
                else:
                    run = p.add_run(part)
                    set_chinese_font(run, "宋体", Pt(12))

        elif elem_type == 'hr':
            # 添加一条分隔线（用段落边框模拟）
            p = doc.add_paragraph()
            p.space_after = Pt(6)

    doc.save(output_path)
    return output_path


def main():
    if len(sys.argv) < 2:
        print("用法：python3 to_docx.py input.md [output.docx]")
        sys.exit(1)

    input_path = Path(sys.argv[1])

    if not input_path.exists():
        print(f"错误：找不到文件 {input_path}")
        sys.exit(1)

    # 确定输出路径
    if len(sys.argv) >= 3:
        output_path = Path(sys.argv[2])
    else:
        output_path = input_path.with_suffix('.docx')

    # 读取 Markdown
    md_text = input_path.read_text(encoding='utf-8')

    # 解析并生成
    elements = parse_markdown(md_text)
    result = create_docx(elements, str(output_path), input_path.parent.resolve())

    print(f"已生成 Word 文档：{result}")


if __name__ == '__main__':
    main()
